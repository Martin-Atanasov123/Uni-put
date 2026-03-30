#!/usr/bin/env python3
import sys, io, json
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document

import glob
matches = glob.glob(r'C:\Users\atana\Downloads\*РУСЕНСКИ*')
path = matches[0]
print(f"File: {path}")
doc = Document(path)

print(f"Paragraphs: {len(doc.paragraphs)}")
print(f"Tables: {len(doc.tables)}")
print()

# Print paragraphs (non-empty)
for i, p in enumerate(doc.paragraphs):
    if p.text.strip():
        print(f"P{i}: {p.text.strip()[:200]}")

print("\n" + "="*80 + "\n")

# Print tables
for ti, table in enumerate(doc.tables):
    print(f"TABLE {ti}: {len(table.rows)} rows x {len(table.columns)} cols")
    for ri, row in enumerate(table.rows):
        cells = []
        for cell in row.cells:
            txt = cell.text.strip().replace('\n', ' | ')
            cells.append(txt[:100] if txt else "NULL")
        print(f"  [{ri}] {json.dumps(cells, ensure_ascii=False)}")
    print()
